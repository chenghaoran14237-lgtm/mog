from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "thesis" / "assets"
FIGURE_DIR = ASSET_DIR / "figures"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
SNAPSHOT_FILE = ASSET_DIR / "project_snapshot.json"
OUTPUT = ROOT / "docs" / "thesis" / "drafts" / "论文正文_程浩然_基于LangGraph的多Agent协作框架在医疗审计场景的设计与实现.docx"
FALLBACK_OUTPUT = ROOT / "docs" / "thesis" / "drafts" / "论文正文_程浩然_自动保存副本.docx"
TOC_PAGES_FILE = ASSET_DIR / "thesis_toc_pages.json"

TITLE = "基于 LangGraph 的多 Agent 协作框架在医疗审计场景的设计与实现"
SCHOOL = "北京信息科技大学"
COLLEGE = "计算机学院"
MAJOR = "计算机科学与技术"
CLASS_ID = "计科2401-本（2024070013）"
STUDENT = "程浩然"
ADVISOR = "乔文豹"
DATE_RANGE = "2026年2月23日至2026年5月22日"

ACCENT = "000000"
LIGHT = "F2F2F2"
GRID = "D7DEE5"
REFERENCE_COUNT = 24
CITATION_FONT_SIZE = 10.5


def load_snapshot() -> dict:
    return json.loads(SNAPSHOT_FILE.read_text(encoding="utf-8"))


def set_run_font(
    run,
    size: float = 10.5,
    *,
    bold: bool = False,
    name: str = "宋体",
    color: str | None = None,
    superscript: bool = False,
) -> None:
    latin_name = "Times New Roman" if name == "Times New Roman" else name
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), latin_name)
    r_fonts.set(qn("w:hAnsi"), latin_name)
    r_fonts.set(qn("w:cs"), latin_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.superscript = superscript


def set_paragraph_format(paragraph, *, first_line: bool = False, align=None, before: float = 0, after: float = 0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = Pt(18)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    if align is not None:
        paragraph.alignment = align


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = Pt(18)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size, bold, outline in [
        ("Heading 1", 16, True, 0),
        ("Heading 2", 12, True, 1),
        ("Heading 3", 12, True, 2),
    ]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(12 if outline == 0 else 6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = Pt(18)
        if style.font.color:
            style.font.color.rgb = RGBColor.from_string("000000")

    clear_header_footer(section)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def clear_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_paragraph(section.header.paragraphs[0])
    clear_paragraph(section.footer.paragraphs[0])


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def set_page_numbering(section, *, start: int | None = None, fmt: str | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    if fmt is not None:
        pg_num_type.set(qn("w:fmt"), fmt)


def configure_section(section, *, header_text: str | None, page_start: int | None, page_fmt: str | None) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    clear_header_footer(section)
    if header_text:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(header, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        run = header.add_run(header_text)
        set_run_font(run, 9, bold=True, name="宋体")
        add_bottom_border(header)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE")
    for run in footer.runs:
        set_run_font(run, 9, name="宋体")
    set_page_numbering(section, start=page_start, fmt=page_fmt)


def add_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder or instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)
    return run


def add_center_text(doc: Document, text: str, size: float, *, bold: bool = False, font_name: str = "宋体", after: float = 0) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=after)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, name=font_name)


def reference_bookmark_name(number: int) -> str:
    return f"ref_{number:03d}"


def citation_numbers(text: str) -> list[int]:
    numbers = [int(part) for part in re.findall(r"\d+", text)]
    if not numbers:
        return []
    if any(number < 1 or number > REFERENCE_COUNT for number in numbers):
        return []
    return numbers


def add_reference_field(paragraph, number: int, *, size: float) -> None:
    # Use Word REF fields so citation numbers are cross-references, not static text.
    run = add_field(paragraph, f"REF {reference_bookmark_name(number)} \\h", str(number))
    set_run_font(run, size, name="Times New Roman", superscript=True)


def add_citation_run(paragraph, citation: str, *, size: float) -> None:
    inner = citation[1:-1]
    if not citation_numbers(inner):
        run = paragraph.add_run(citation)
        set_run_font(run, CITATION_FONT_SIZE, name="Times New Roman", superscript=True)
        return
    run = paragraph.add_run("[")
    set_run_font(run, CITATION_FONT_SIZE, name="Times New Roman", superscript=True)
    for part in re.split(r"(\d+)", inner):
        if not part:
            continue
        if part.isdigit():
            add_reference_field(paragraph, int(part), size=CITATION_FONT_SIZE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, CITATION_FONT_SIZE, name="Times New Roman", superscript=True)
    run = paragraph.add_run("]")
    set_run_font(run, CITATION_FONT_SIZE, name="Times New Roman", superscript=True)


def add_bookmarked_reference_number(paragraph, number: int, *, size: float) -> None:
    bookmark_id = 1000 + number
    run = paragraph.add_run(str(number))
    set_run_font(run, size, name="Times New Roman")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), reference_bookmark_name(number))
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    parent = paragraph._p
    run_index = parent.index(run._r)
    parent.insert(run_index, start)
    parent.insert(run_index + 2, end)


def add_runs_with_citations(paragraph, text: str, *, bold: bool = False, size: float = 10.5) -> None:
    for part in re.split(r"(\[[0-9,\-\s]+\])", text):
        if not part:
            continue
        if re.fullmatch(r"\[[0-9,\-\s]+\]", part):
            add_citation_run(paragraph, part, size=size)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, bold=bold)


def add_body(doc: Document, text: str, *, first_line: bool = True, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, 10.5, bold=True)
        add_runs_with_citations(p, text[len(bold_prefix) :])
    else:
        add_runs_with_citations(p, text)


def add_chapter(doc: Document, title: str) -> None:
    if title.startswith(("第二章", "第三章", "第四章", "第五章", "结束语", "参考文献", "致谢", "附录")):
        doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 16, bold=True, name="黑体", color="000000")


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(title)
    set_run_font(run, 12, bold=True, name="黑体", color="000000")


def add_subsection(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Heading 3")
    run = p.add_run(title)
    set_run_font(run, 11, bold=True, name="黑体", color="000000")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, header: bool = False, size: float = 9.5) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=0)
    if header:
        run = p.add_run(str(text))
        set_run_font(run, size, bold=True, name="宋体")
    else:
        add_runs_with_citations(p, str(text), size=size)
    if header:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, "FFFFFF")


def add_table_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    run = p.add_run(caption)
    set_run_font(run, 10.5, bold=True)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, header=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "FFFFFF")
        borders.append(border)
    tbl_pr.append(borders)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, *, left: int = 0, right: int = 0, top: int = 0, bottom: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "left": left, "bottom": bottom, "right": right}.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_figure(doc: Document, filename: str, caption: str, *, width_cm: float = 14.8) -> None:
    path = FIGURE_DIR / filename
    add_picture_with_caption(doc, path, caption, width_cm=width_cm)


def add_screenshot(doc: Document, filename: str, caption: str, *, width_cm: float = 14.8) -> None:
    path = SCREENSHOT_DIR / filename
    add_picture_with_caption(doc, path, caption, width_cm=width_cm)


def add_picture_with_caption(doc: Document, path: Path, caption: str, *, width_cm: float = 14.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Body paragraphs use a fixed 18 pt line grid; image paragraphs must override
    # it or LibreOffice clips inline pictures to a single text line.
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    set_paragraph_format(cap, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=6)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, 10.5, bold=True)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        add_cover_blank(doc, centered=True)
    add_center_text(doc, SCHOOL, 26, bold=False, font_name="楷体")
    add_cover_blank(doc, centered=True)
    add_center_text(doc, "毕业设计（论文）", 42, bold=True, font_name="楷体")
    for _ in range(8):
        add_cover_blank(doc)
    add_cover_form(doc)


def add_cover_blank(doc: Document, *, centered: bool = False, first_line_indent: float | None = None) -> None:
    p = doc.add_paragraph()
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)


def split_cover_title(title: str) -> list[str]:
    marker = "在医疗审计"
    if marker in title:
        left, right = title.split(marker, 1)
        return [left + marker, right]
    return [title]


def clear_cell(cell) -> None:
    cell.text = ""
    for paragraph in cell.paragraphs:
        clear_paragraph(paragraph)


def set_cover_cell_text(cell, text: str, *, bold: bool = False, underline: bool = False, align=None, size: float = 14) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, left=40, right=40, top=0, bottom=0)
    clear_cell(cell)
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, name="宋体")
    run.underline = underline


def set_cover_row_widths(row) -> None:
    for cell, width in zip(row.cells, [3.0, 2.2, 3.0, 7.0], strict=True):
        set_cell_width(cell, width)


def add_cover_form_row(table, label: str, value: str) -> None:
    row = table.add_row()
    set_cover_row_widths(row)
    value_cell = row.cells[1].merge(row.cells[3])
    set_cover_cell_text(row.cells[0], label, bold=True)
    set_cover_cell_text(value_cell, value, underline=True)


def add_cover_form(doc: Document) -> None:
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    title_row = table.add_row()
    set_cover_row_widths(title_row)
    title_cell = title_row.cells[1].merge(title_row.cells[3])
    set_cover_cell_text(title_row.cells[0], "题    目：", bold=True)
    clear_cell(title_cell)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(title_cell, left=40, right=40, top=0, bottom=0)
    for index, line in enumerate(split_cover_title(TITLE)):
        p = title_cell.paragraphs[0] if index == 0 else title_cell.add_paragraph()
        set_paragraph_format(p, after=0)
        run = p.add_run(line)
        set_run_font(run, 14, name="宋体")
        run.underline = True

    add_cover_form_row(table, "学    院：", COLLEGE)
    add_cover_form_row(table, "专    业：", MAJOR)
    add_cover_form_row(table, "学生姓名：", STUDENT)
    add_cover_form_row(table, "班级/学号：", CLASS_ID.replace("（", "/").replace("）", ""))

    row = table.add_row()
    set_cover_row_widths(row)
    label_cell = row.cells[0].merge(row.cells[1])
    value_cell = row.cells[2].merge(row.cells[3])
    set_cover_cell_text(label_cell, "指导老师/督导老师：", bold=True)
    set_cover_cell_text(value_cell, ADVISOR, underline=True)

    add_cover_form_row(table, "起止时间：", DATE_RANGE)


def add_cover_line(
    doc: Document,
    label: str,
    value: str,
    *,
    width_chars: int,
    trailing_blanks: int = 2,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(90)
    label_run = p.add_run(label)
    set_run_font(label_run, 14, bold=True, name="宋体")
    value_run = p.add_run(value.ljust(width_chars))
    set_run_font(value_run, 14, name="宋体")
    value_run.underline = True
    for _ in range(trailing_blanks):
        add_cover_blank(doc, first_line_indent=90)


def add_cover_student_line(doc: Document, student: str, class_id: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(90)
    label_run = p.add_run("学生姓名：")
    set_run_font(label_run, 14, bold=True, name="宋体")
    student_run = p.add_run(student.ljust(10))
    set_run_font(student_run, 14, name="宋体")
    student_run.underline = True
    class_label_run = p.add_run("   班级/学号")
    set_run_font(class_label_run, 14, bold=True, name="宋体")
    class_run = p.add_run(("  " + class_id).ljust(24))
    set_run_font(class_run, 14, name="宋体")
    class_run.underline = True
    for _ in range(2):
        add_cover_blank(doc, first_line_indent=90)


def add_front_matter(doc: Document) -> None:
    add_chapter(doc, "毕业设计（论文）任务书")
    task_rows = [
        ["题目", TITLE],
        ["学生", f"{STUDENT}，{CLASS_ID}"],
        ["指导教师", ADVISOR],
        ["主要任务", "实现面向个人健康资料的医疗审计辅助系统，完成资料接入、OCR、标准化入库、指标查询、LangGraph 多 Agent 审计报告生成和前端展示。"],
        ["主要成果", "可运行系统、数据库与接口说明、测试记录、论文初稿及项目图表。"],
    ]
    add_table(doc, "表0-1 毕业设计任务摘要", ["项目", "内容"], task_rows)
    add_body(doc, "本任务书根据开题报告确定的实现路线整理，系统以“资料接入、OCR 识别、结构化抽取、版本化存储、指标查询、多 Agent 审计、结果展示”为主线，避免将系统定位为临床诊断工具，而是定位为医疗资料质量审计、风险提示与证据追溯的工程实现。")
    doc.add_page_break()

    add_chapter(doc, "毕业设计（论文）原创性声明")
    add_body(doc, f"本人郑重声明：所呈交的毕业设计（论文），题目为《{TITLE}》，是在指导教师指导下独立完成的设计与实现成果。除文中特别标注和引用的内容外，本文不包含他人已经发表或撰写过的成果，也不包含为获得北京信息科技大学或其他教育机构学位、证书而使用过的材料。")
    add_body(doc, "本人承诺文中所用图表、接口、字段、测试数据和系统流程均来自本项目源码、数据库模型、运行脚本或公开文献，引用资料已在参考文献中列出。")
    add_body(doc, "作者签名：__________________    年    月    日")
    doc.add_page_break()

    add_chapter(doc, "毕业设计（论文）版权授权声明")
    add_body(doc, "本人完全了解北京信息科技大学关于收集、保存和使用本科毕业设计（论文）的有关要求，同意学校保留并向有关机构提交论文的电子版和纸质版，允许论文被查阅和借阅。本人授权学校采用影印、缩印、数字化或其他复制方式保存论文。")
    add_body(doc, "作者签名：__________________    指导教师签名：__________________    年    月    日")
    doc.add_page_break()


def add_abstracts(doc: Document) -> None:
    add_chapter(doc, "摘   要")
    abstract = (
        "针对个人体检报告、检验单和病历摘要等医疗资料来源分散、格式不统一、处理过程缺少追溯的问题，本文设计并实现了一套基于 LangGraph 的多 Agent 协作医疗审计辅助系统。"
        "系统采用前后端分离架构，后端以 FastAPI、SQLAlchemy 和 MySQL 为基础，围绕原始文件、OCR 结果、文档版本、结构化指标、任务事件和审计事件建立数据模型；"
        "同时通过 Provider 抽象接入 OCR、标准化、大语言模型和 RAG 知识检索能力。综合审计报告模块以 LangGraph 状态图组织文档质量审计、时间线构建、指标一致性检查、风险提示、知识检索、证据绑定、冲突复核、合规检查、报告生成和安全审查等节点，"
        "通过条件路由和回环机制保证审计结论能够补充证据、回退复核并持久化。前端实现文档接入、文档库、综合审计报告、智能洞察、RAG 知识库、指标探索和任务状态监控，能够直观看到审计节点与边的流转。测试结果表明，系统能够基于 Mock 体检数据完成从资料入库到报告生成的闭环。"
    )
    add_body(doc, abstract)
    add_body(doc, "关键词：LangGraph；多 Agent；医疗审计；OCR；结构化入库；证据追溯；FastAPI", first_line=False, bold_prefix="关键词：")
    doc.add_page_break()

    add_chapter(doc, "Abstract")
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=True)
    run = p.add_run(
        "This thesis designs and implements a LangGraph-based multi-agent collaboration framework for medical audit scenarios. "
        "The system targets personal health documents such as physical examination reports, laboratory reports and clinical notes. "
        "It builds a traceable data pipeline covering upload, OCR, normalization, versioned storage, metric querying, audit orchestration and frontend visualization. "
        "The backend is implemented with FastAPI, SQLAlchemy and MySQL, and the external OCR, normalization and LLM capabilities are accessed through replaceable providers. "
        "The audit report module uses LangGraph as a state machine and coordinates document quality checking, timeline construction, metric consistency review, risk analysis, RAG-based knowledge retrieval, evidence binding, conflict review, compliance checking, report composition and safety review. "
        "Conditional routing and cyclic transitions make the workflow auditable and suitable for iterative evidence completion. "
        "The frontend presents document ingestion, document management, RAG knowledge search, task monitoring, metric exploration, intelligent insight and an audit report workspace with visible node and edge transitions. "
        "Tests based on mock health examination data show that the system can complete an end-to-end workflow from document ingestion to final audit report generation."
    )
    set_run_font(run, 12, name="Times New Roman")
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    run = p.add_run("Keywords: LangGraph; multi-agent; medical audit; OCR; normalization; evidence traceability; FastAPI")
    set_run_font(run, 12, bold=True, name="Times New Roman")
    doc.add_page_break()


def add_chinese_abstract(doc: Document) -> None:
    add_chapter(doc, "摘    要")
    abstract = (
        "针对个人体检报告、检验单和病历摘要等医疗资料来源分散、格式不统一、处理过程缺少追溯的问题，本文设计并实现了一套基于 LangGraph 的多 Agent 协作医疗审计辅助系统。"
        "系统采用前后端分离架构，后端以 FastAPI、SQLAlchemy 和 MySQL 为基础，围绕原始文件、OCR 结果、文档版本、结构化指标、任务事件和审计事件建立数据模型；"
        "同时通过 Provider 抽象接入 OCR、标准化、大语言模型和 RAG 知识检索能力。综合审计报告模块以 LangGraph 状态图组织文档质量审计、时间线构建、指标一致性检查、风险提示、知识检索、证据绑定、冲突复核、合规检查、报告生成和安全审查等节点，"
        "通过条件路由和回环机制保证审计结论能够补充证据、回退复核并持久化。前端实现文档接入、文档库、综合审计报告、智能洞察、RAG 知识库、指标探索和任务状态监控，能够直观看到审计节点与边的流转。测试结果表明，系统能够基于 Mock 体检数据完成从资料入库到报告生成的闭环。"
    )
    add_body(doc, abstract)
    add_body(doc, "关键词：LangGraph；多 Agent；医疗审计；OCR；结构化入库；证据追溯；FastAPI", first_line=False, bold_prefix="关键词：")


def add_english_abstract(doc: Document) -> None:
    add_chapter(doc, "Abstract")
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=True)
    run = p.add_run(
        "This thesis designs and implements a LangGraph-based multi-agent collaboration framework for medical audit scenarios. "
        "The system targets personal health documents such as physical examination reports, laboratory reports and clinical notes. "
        "It builds a traceable data pipeline covering upload, OCR, normalization, versioned storage, metric querying, audit orchestration and frontend visualization. "
        "The backend is implemented with FastAPI, SQLAlchemy and MySQL, and the external OCR, normalization and LLM capabilities are accessed through replaceable providers. "
        "The audit report module uses LangGraph as a state machine and coordinates document quality checking, timeline construction, metric consistency review, risk analysis, RAG-based knowledge retrieval, evidence binding, conflict review, compliance checking, report composition and safety review. "
        "Conditional routing and cyclic transitions make the workflow auditable and suitable for iterative evidence completion. "
        "The frontend presents document ingestion, document management, RAG knowledge search, task monitoring, metric exploration, intelligent insight and an audit report workspace with visible node and edge transitions. "
        "Tests based on mock health examination data show that the system can complete an end-to-end workflow from document ingestion to final audit report generation."
    )
    set_run_font(run, 12, name="Times New Roman")
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    run = p.add_run("Keywords: LangGraph; multi-agent; medical audit; OCR; normalization; evidence traceability; FastAPI")
    set_run_font(run, 12, bold=True, name="Times New Roman")


def load_toc_pages() -> dict[str, str]:
    defaults = {
        "I": "I",
        "II": "II",
        "第一章 绪论": "1",
        "1.1 实现背景和意义": "1",
        "1.2 国内外实现现状": "2",
        "1.3 本文主要工作": "3",
        "第二章 相关技术与系统基础": "5",
        "2.1 医疗资料标准化与电子病历文本挖掘": "6",
        "2.2 医疗大模型、知识增强与 RAG 方法": "6",
        "2.3 Agent 工作流与状态机编排": "7",
        "2.4 Provider 抽象与外部能力接入": "8",
        "2.5 本项目已有实现基础": "9",
        "第三章 需求分析与总体设计": "12",
        "3.1 需求分析": "12",
        "3.2 系统总体架构": "13",
        "3.3 业务流程设计": "13",
        "3.4 后端分层与接口设计": "14",
        "3.5 数据库设计": "16",
        "3.6 安全与可追溯设计": "18",
        "第四章 系统详细实现": "21",
        "4.1 文件上传与 OCR 处理实现": "21",
        "4.2 标准化与版本化入库实现": "21",
        "4.3 综合审计报告 LangGraph 实现": "22",
        "4.4 审计事件与节点状态持久化": "24",
        "4.5 前端模块实现": "25",
        "第五章 测试与运行效果分析": "28",
        "5.1 Mock 体检数据构造": "28",
        "5.2 接口与端到端测试": "29",
        "5.3 真实演示运行截图": "31",
        "5.4 运行效果与不足分析": "34",
        "结束语": "36",
        "参考文献": "37",
        "致谢": "39",
    }
    if TOC_PAGES_FILE.exists():
        defaults.update(json.loads(TOC_PAGES_FILE.read_text(encoding="utf-8")))
    return defaults


def set_toc_cell(cell, text: str, *, level: int = 0, bold: bool = False, align=None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, first_line=False, align=align, after=0)
    p.paragraph_format.left_indent = Cm(0.65 * level)
    run = p.add_run(text)
    set_run_font(run, 10.5 if not bold else 11, bold=bold, name="宋体")


def add_toc_line(table, row_index: int, title: str, page: str, *, level: int = 0, bold: bool = False) -> None:
    row = table.rows[row_index]
    set_cell_width(row.cells[0], 14.2)
    set_cell_width(row.cells[1], 1.2)
    set_toc_cell(row.cells[0], title, level=level, bold=bold)
    set_toc_cell(row.cells[1], page, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT)


def add_toc(doc: Document) -> None:
    pages = load_toc_pages()
    add_chapter(doc, "目    录")
    entries = [
        ("摘要（中文）", "I", 0, True),
        ("（英文）", "II", 1, False),
        ("第一章 绪论", "第一章 绪论", 0, True),
        ("1.1 实现背景和意义", "1.1 实现背景和意义", 1, False),
        ("1.2 国内外实现现状", "1.2 国内外实现现状", 1, False),
        ("1.3 本文主要工作", "1.3 本文主要工作", 1, False),
        ("第二章 相关技术与系统基础", "第二章 相关技术与系统基础", 0, True),
        ("2.1 医疗资料标准化与电子病历文本挖掘", "2.1 医疗资料标准化与电子病历文本挖掘", 1, False),
        ("2.2 医疗大模型、知识增强与 RAG 方法", "2.2 医疗大模型、知识增强与 RAG 方法", 1, False),
        ("2.3 Agent 工作流与状态机编排", "2.3 Agent 工作流与状态机编排", 1, False),
        ("2.4 Provider 抽象与外部能力接入", "2.4 Provider 抽象与外部能力接入", 1, False),
        ("2.5 本项目已有实现基础", "2.5 本项目已有实现基础", 1, False),
        ("第三章 需求分析与总体设计", "第三章 需求分析与总体设计", 0, True),
        ("3.1 需求分析", "3.1 需求分析", 1, False),
        ("3.2 系统总体架构", "3.2 系统总体架构", 1, False),
        ("3.3 业务流程设计", "3.3 业务流程设计", 1, False),
        ("3.4 后端分层与接口设计", "3.4 后端分层与接口设计", 1, False),
        ("3.5 数据库设计", "3.5 数据库设计", 1, False),
        ("3.6 安全与可追溯设计", "3.6 安全与可追溯设计", 1, False),
        ("第四章 系统详细实现", "第四章 系统详细实现", 0, True),
        ("4.1 文件上传与 OCR 处理实现", "4.1 文件上传与 OCR 处理实现", 1, False),
        ("4.2 标准化与版本化入库实现", "4.2 标准化与版本化入库实现", 1, False),
        ("4.3 综合审计报告 LangGraph 实现", "4.3 综合审计报告 LangGraph 实现", 1, False),
        ("4.4 审计事件与节点状态持久化", "4.4 审计事件与节点状态持久化", 1, False),
        ("4.5 前端模块实现", "4.5 前端模块实现", 1, False),
        ("第五章 测试与运行效果分析", "第五章 测试与运行效果分析", 0, True),
        ("5.1 Mock 体检数据构造", "5.1 Mock 体检数据构造", 1, False),
        ("5.2 接口与端到端测试", "5.2 接口与端到端测试", 1, False),
        ("5.3 真实演示运行截图", "5.3 真实演示运行截图", 1, False),
        ("5.4 运行效果与不足分析", "5.4 运行效果与不足分析", 1, False),
        ("结束语", "结束语", 0, True),
        ("参考文献", "参考文献", 0, True),
        ("致谢", "致谢", 0, True),
    ]
    table = doc.add_table(rows=len(entries), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    for idx, (title, key, level, bold) in enumerate(entries):
        add_toc_line(table, idx, title, pages.get(key, ""), level=level, bold=bold)


EXPANSION_TEXT: dict[str, list[str]] = {
    "chapter_one": [
        "从本科毕业设计的工程定位看，本文更强调“设计与实现”而不是纯理论研究。医疗审计场景中的核心难点并不只是调用一次大语言模型，而是如何把输入资料、结构化数据、审计过程、证据来源和最终报告串成一个能够复现的系统。若系统只在前端提供一个聊天框，用户虽然可以得到自然语言回答，但无法知道回答依据来自哪一份报告、哪一个指标、哪一次 OCR 或哪一个审计节点。因此本文在课题实现中把数据结构、服务接口和状态机执行过程放在同等重要的位置，保证系统不仅能够生成文本，还能够解释文本生成过程。",
        "本文所说的医疗审计不是医疗诊断，也不是替代医生给出治疗建议，而是面向个人健康资料管理场景，对多来源资料进行一致性检查、异常指标提示、证据绑定和报告汇总。这样的边界设计符合本科工程项目的可实现范围：一方面，系统能够展示 OCR、标准化、数据库建模、后端接口、前端交互和多 Agent 状态图等技术点；另一方面，系统不会越过临床诊断边界，避免把毕业设计做成无法验证的医疗问答系统。论文后续章节中的数据模型、流程图和测试用例都围绕这一边界展开。",
        "在资料输入方面，普通用户能够提供的材料往往不是统一格式。体检报告中可能包含总检结论和多个科室检查结果，化验报告中可能包含指标、单位、参考范围和异常标记，病历摘要中则可能以自然语言描述既往史、用药史和建议事项。若直接将这些内容作为长文本交给模型处理，模型输出很难稳定复现，也难以回到原始字段。本文采用先 OCR、再标准化、再审计的实现路线，是为了把非结构化资料逐步转化为可查询、可比较和可追溯的数据对象。",
        "选择 LangGraph 的原因也来自该场景本身。医疗资料审计并不适合只按固定顺序执行一次，因为审计过程可能在多个节点之间回退。例如风险节点发现某项指标异常但证据不足时，需要重新进入证据绑定节点；引用检查发现报告草稿中存在未绑定来源的结论时，需要回到审计阶段补充证据；安全审查发现报告语言接近诊断承诺时，需要返回报告生成节点重写。普通线性工作流只能把这些情况写成大量 if-else，而状态图能够把这些回环作为边关系显式表达出来。",
        "因此，本文的主要目标可以概括为三个层次。第一，完成医疗资料处理底座，使原始文件、OCR 文本、结构化文档、版本快照和指标表形成稳定的数据链路。第二，完成基于 LangGraph 的综合审计报告状态机，使多个 Agent 节点能够围绕同一份 GraphState 协作执行。第三，完成前端可视化展示，使用户能够看到真实事件如何驱动节点高亮和边流转，而不是只看到静态流程图。这三个层次共同构成本文工程实现的核心。",
    ],
    "chapter_two": [
        "医疗资料标准化的关键是确定哪些信息必须结构化，哪些信息可以作为叙事事实保存。对于数值型检验指标，系统需要保存指标名称、数值、单位、参考范围、异常标记和观察时间；对于病历摘要、影像结论和总检建议，系统更适合保存为 prose_facts 或 narrative_context，避免把自然语言强行拆成错误的数值字段。这样的区分能够降低标准化失败对后续审计的影响，也使系统在面对不同报告模板时具有更好的容错能力[4-8]。",
        "在本文系统中，标准化不是一次性覆盖操作，而是一个可版本化过程。每次 OCR 结果进入标准化服务后，系统会生成 ExtractedDocument 当前投影，同时创建 DocumentVersion 保存版本快照。这样做的原因是 OCR Provider、LLM Provider 和规则库都可能发生变化，若系统直接覆盖旧结果，就无法解释历史报告为何得出当时的结论。版本化设计让每次综合审计报告都可以绑定具体 document_version_id，从而保证报告证据能够回到当时使用的结构化结果[8,15]。",
        "Provider 抽象在本项目中承担了工程隔离作用。OCR、LLM、标准化和存储都可能在不同环境下切换具体实现，例如文本样本使用 plaintext，图片样本在密钥配置后使用百度 OCR，部署时还可以接入视觉模型或对象存储；标准化在网络不可用时使用规则兜底，在可用时调用兼容 OpenAI 协议的模型。若业务服务直接依赖某个具体 SDK，后续替换成本会很高。本文通过 ProviderGateway 记录调用事件、耗时、状态和错误信息，使外部能力既可替换，又可被审计[14]。",
        "LangGraph 状态机的核心不是“多写几个函数”，而是把节点之间的控制权交给状态和条件边。AuditGraphState 中保存 selected_document_version_ids、documents、measurements、knowledge_chunks、knowledge_context、completed_agents、route_history、evidence_items、citation_issues、safety_issues、report_draft 和 final_report 等字段。每个节点只负责读取必要字段并返回局部更新，路由节点根据状态决定下一步。这样可以避免单个服务函数无限膨胀，也便于前端根据事件流展示节点执行过程[22]。",
        "多 Agent 协作在本文中采用职责拆分，而不是多个模型互相聊天。document_quality_agent 检查资料是否完整，timeline_builder 组织时间线，measurement_consistency_agent 检查指标变化和异常，risk_agent 生成非诊断性关注点，knowledge_retrieval_agent 检索医学审计知识，evidence_agent 负责证据绑定，conflict_agent 检查叙事事实与结构化指标之间的冲突，compliance_agent 检查结论是否缺少证据，report_composer 负责报告组织，citation_checker 和 safety_reviewer 分别进行引用和安全审查。各节点职责边界清晰，便于测试和论文说明[20-22]。",
        "项目已有基础包括 FastAPI 路由、SQLAlchemy 模型、Pydantic 数据校验、React 前端模块和 Mock 体检数据脚本。这些基础并不是论文外的附属内容，而是支撑 LangGraph 架构落地的工程条件。若没有稳定的数据模型，状态机无法获得可靠输入；若没有事件表和节点状态表，前端无法展示真实流转；若没有 Mock 数据，测试和答辩演示无法稳定复现。因此本章所述基础技术与后续系统实现之间存在直接对应关系[14-15]。",
    ],
    "chapter_three": [
        "需求分析阶段首先需要明确用户角色和使用边界。系统面向普通个人用户和答辩演示场景，用户可以上传体检、化验、病历摘要和影像结论等资料，系统负责将资料转化为结构化记录，并在用户主动选择若干文档后生成综合审计报告。系统不提供诊断结论、不推荐具体治疗方案，也不对医生意见作出替代判断。所有风险提示都应以“关注点”“建议复核”“建议结合医生意见”形式表达，并且必须绑定原始证据。",
        "功能需求可以分为资料接入、资料处理、资料查询、知识检索、智能交互、任务监控和综合审计七类。资料接入负责文件上传和元数据保存；资料处理负责 OCR、标准化、版本快照和指标入库；资料查询负责文档列表、文档详情和指标检索；知识检索负责 RAG 知识来源和命中结果展示；智能交互保留聊天式问答能力，但不承担 LangGraph 审计主流程；任务监控负责 OCR 与标准化状态；综合审计则通过状态机生成最终报告。",
        "非功能需求主要包括可追溯性、可复现性、可扩展性和安全边界。可追溯性要求每个结论都能回到文档版本、指标或 OCR 原文；可复现性要求 Mock 数据和测试账户能够稳定重建演示环境；可扩展性要求 Provider 能力和审计节点可以替换或增加；安全边界要求系统输出避免诊断承诺，并且所有用户数据按 user_id 隔离。上述非功能需求决定了系统不能只做简单 CRUD，而必须在数据库和服务层设计审计事件、节点状态和 Provider 调用记录。",
        "总体架构采用前后端分离设计。前端负责资料选择、模块切换、流程图展示、知识检索、任务状态和报告阅读；后端提供认证、文件、OCR、标准化、文档、指标、任务、知识库和审计报告等接口；数据库保存用户数据、任务事件、Provider 调用事件、知识块和审计运行记录。综合审计报告模块启动后，后端创建 AuditReportRun，并在执行过程中持续写入 AuditReportEvent 和 AuditReportNodeState，前端通过轮询接口读取这些事件，从而把后端真实状态反映到页面上。",
        "数据库设计围绕三条主链路展开。第一条是资料处理链路，从 users 到 records、record_files、ocr_results、extracted_documents、document_versions 和 measurements，负责把原始资料转化为结构化指标。第二条是异步任务链路，从 tasks 到 task_events 和 provider_events，负责记录后台处理过程和外部能力调用。第三条是审计报告链路，从 audit_report_runs 到 audit_report_events 和 audit_report_node_states，负责记录 LangGraph 执行过程。三条链路分别对应论文 ER 图中的蓝色、橙色和紫色区域。",
        "接口设计遵循资源边界清晰的原则。认证接口只负责登录和当前用户信息，文件接口只负责上传和文件元数据，OCR 接口只负责抽取任务，ingestion 接口只负责标准化入库，documents 和 measurements 接口负责查询结构化结果，audit-reports 接口负责创建、执行、轮询和读取综合报告。这样的划分使前端模块可以按业务目标组织页面，也使后端服务在测试时能够逐段验证，不需要通过一个复杂接口完成所有动作。",
        "综合审计报告的状态机设计是本章重点。audit_router 是主路由节点，它根据 completed_agents 和当前状态判断下一个审计节点；普通 Agent 节点执行完毕后不直接进入下一个 Agent，而是回到 audit_router，由路由节点再次判断是否需要继续分派、补充证据或进入报告生成。final_router 是终态路由节点，它在引用检查和安全审查之后决定是否持久化报告，或者回到审计阶段和报告生成阶段。这种双路由结构体现了状态机的必要性。",
        "前端设计上，综合审计报告模块不再复用智能洞察的聊天布局，而是独立成为一个以流程图为核心的工作台。左侧只保留必要的文档选择，主体区域展示 LangGraph 流转图，底部或侧边提供最终报告入口。节点高亮不由前端定时模拟，而是来自后端事件数据：event_type、node_name、edge_source、edge_target 和 status 决定节点和边的显示状态。这样用户看到的图既有展示效果，也能对应数据库中的真实事件。",
    ],
    "chapter_four": [
        "文件上传实现中，后端首先验证用户身份，然后创建 records 和 record_files。record_files 不仅保存文件名和类型，也保存 storage_provider、storage_key、size_bytes 和 content_bytes 等字段。当前演示环境使用 database_inline 存储，是为了降低部署复杂度并保证测试可复现；若后续接入对象存储，业务层只需替换 StorageProvider，不需要修改 OCR 和标准化主流程。这体现了 Provider 抽象在工程实现中的实际价值。",
        "OCR 处理并不是简单返回字符串。系统会为 OCR 创建后台任务，记录 tasks.status，并在 task_events 中保存状态变化。OCRProvider 执行后，系统将 raw_text、raw_payload、provider_name、revision_number 和 is_current 写入 ocr_results。若同一文件重新识别，旧 OCR 结果不会立即删除，而是通过 revision_number 和 supersedes_ocr_result_id 保留历史关系。这样可以解释不同时间生成的标准化结果为何存在差异。",
        "标准化服务读取 OCRResult.raw_text 后调用 NormalizationProvider，得到 document_type、document_category、report_date、measurements 和 prose_facts。对于可数值化的指标，系统写入 measurements；对于病历描述和影像结论，系统保存到 normalized_payload 中。为提升稳定性，服务在 LLM 输出缺字段或解析失败时使用规则兜底，至少保证日期、常见指标和叙事事实能够被提取。这样的处理策略比完全依赖模型输出更适合毕业设计演示和自动化测试。",
        "LangGraph 审计引擎的实现分为状态初始化、节点执行、路由判断、事件输出和结果持久化几个部分。load_graph_state 从数据库读取用户选择的 document_versions 和关联 measurements，构造 AuditGraphState。每个 Agent 节点只负责一个有限任务，并把结果写入对应字段。Engine.stream 在每次节点执行后产出事件，服务层立即写入 audit_report_events，使前端可以在运行过程中轮询到节点变化，而不是等最终报告完成后一次性返回。",
        "audit_router 的路由策略以 completed_agents 为核心。若文档质量节点尚未执行，则先进入 document_quality_agent；若时间线尚未构建，则进入 timeline_builder；若指标一致性、风险、知识检索、证据、冲突和合规节点尚未完成，则依次进入对应 Agent；当必要审计节点完成后，进入 quality_gate。quality_gate 不负责生成报告，而是判断当前状态是否具备进入 report_composer 的最低条件。这个设计避免报告生成节点承担过多判断职责。",
        "final_router 的价值在于处理报告生成后的回环。report_composer 生成草稿后，citation_checker 会检查报告中的关键结论是否绑定 evidence_items；safety_reviewer 会检查报告语言是否越过医疗审计边界。若 citation_issues 存在，final_router 将流程引回 audit_router 补充审计；若 safety_issues 存在，则引回 report_composer 重写报告；只有当引用和安全审查通过，或者达到最大迭代保护条件时，流程才进入 persist_report。该机制体现了 LangGraph 与线性工作流的区别。",
        "审计事件持久化是前端流程图真实流动的基础。AuditReportEvent 保存 sequence、event_type、node_name、edge_source、edge_target、status、message 和 payload；AuditReportNodeState 保存 node_name、status、visit_count、last_event_id 和 output。当前端轮询 events 和 nodes 接口时，可以根据 edge_source 和 edge_target 高亮边，根据 node_name 和 status 高亮节点，根据 visit_count 展示节点是否被回环访问。这样，前端动画不是装饰，而是后端状态的可视化映射。",
        "综合审计报告的最终结构包括资料概览、时间线、指标异常、风险关注点、知识来源、证据列表、冲突提示、合规说明和非诊断性声明。报告中的每一条关键结论都应尽量绑定 evidence_items，证据对象包含 document_version_id、measurement_id、source_text、knowledge_chunk 或字段路径。若某条结论无法绑定证据，citation_checker 会把它标记为问题。这样的报告结构比普通自然语言总结更适合医疗资料审计，因为它强调来源、依据和复核边界。",
        "前端实现中，综合审计报告模块、智能洞察模块和 RAG 知识库保持分离。智能洞察保持 chatbot 形态并显示历史记录；RAG 知识库展示知识块来源和检索命中；综合审计报告则以流程图和最终报告为核心。用户选择文档后启动运行，前端周期性请求事件和节点状态，更新图中节点、边和报告按钮。任务状态按钮集中展示 OCR 与标准化处理状态，使长耗时流程具有可观察性。",
        "异常处理方面，系统需要处理空文档、无 OCR 文本、标准化失败、缺少 report_date、无 measurements、Provider 超时和 LLM 输出格式错误等情况。当前实现采用任务状态、错误消息和规则兜底降低失败影响。对于综合审计报告，若输入文档不足，document_quality_agent 会输出质量发现；若证据不足，evidence_agent 和 citation_checker 会触发回环；若迭代次数过多，final_router 会根据 max_iterations 保护流程结束，避免无限循环。",
    ],
    "chapter_five": [
        "Mock 数据的作用不是为了凑数，而是为了覆盖端到端流程。当前数据集中包含体检报告、化验报告、病历摘要和影像结论，既有结构化指标，也有自然语言事实。综合审计报告选择多份文档运行时，系统可以同时测试时间线构建、指标异常识别、证据绑定和报告生成。如果只有一两份简单文档，LangGraph 中的多个 Agent 节点很难体现差异，也无法证明回环和状态持久化的必要性。",
        "接口测试重点覆盖主链路而不是孤立函数。登录接口验证认证能力，文件上传接口验证 records 和 record_files 写入，OCR 接口验证任务和 ocr_results 写入，标准化接口验证 ExtractedDocument、DocumentVersion 和 Measurement 写入，文档与指标查询接口验证结构化结果可读，综合审计接口验证 AuditReportRun、AuditReportEvent、AuditReportNodeState 和 final_report 是否完整生成。该测试顺序与用户真实操作顺序一致。",
        "对于 LangGraph 审计报告，测试不能只检查最终报告是否存在，还要检查事件和节点状态是否符合预期。至少需要确认 audit_router 被访问多次，普通 Agent 节点执行后回到 audit_router，report_composer 之后进入 citation_checker 和 safety_reviewer，final_router 最终进入 persist_report。若事件表中只有一条直线式记录，则说明系统没有真正体现状态机；若 visit_count 和 route_history 能反映回环，则可以支撑论文和答辩中的技术说明。",
        "运行效果分析应从系统闭环、技术体现和不足三个角度说明。从系统闭环看，用户可以从上传资料开始，经过 OCR、标准化和指标查询，最终生成综合审计报告；从技术体现看，LangGraph 状态机通过节点、边、状态字段、事件表和前端流程图形成可见链路；从不足看，真实图片 OCR、复杂表格识别、指标单位换算、参考范围解析和 LLM 输出稳定性仍需要继续加强。这种分析比简单写“测试通过”更符合毕业论文要求。",
        "当前版本的主要优势是演示路径清晰。答辩时可以先展示 Mock 账户中的多份体检资料，再展示文档标准化结果和指标搜索，随后选择若干文档启动综合审计报告。运行过程中，前端流程图会根据后端事件高亮节点和边，最终通过按钮打开完整报告。该演示能够直接回答中期反馈中“没有体现 LangGraph 架构”的问题，因为用户可以看到状态图如何驱动真实数据流转。",
        "当前版本的主要不足也需要在论文中如实说明。第一，OCR 在本地主要使用 plaintext，真实图片和扫描 PDF 的识别质量需要接入真实 OCR 服务后继续验证。第二，标准化仍依赖提示词、规则兜底和常见指标别名，对复杂报告模板的适配还不充分。第三，审计 Agent 以规则和结构化检查为主，部分语言解释能力仍可增强。第四，前端流程图虽然已经能够展示真实流转，但在不同屏幕尺寸下仍需要进一步优化画布缩放和节点密度。",
        "后续完善可以分为短期和长期两类。短期工作包括补充更多异常测试、完善数据库迁移脚本、增加图片型样本、优化指标别名表和单位换算表，并在论文定稿前核对所有图表题注和参考文献格式。长期工作包括接入真实 OCR 服务、引入更稳定的结构化抽取模型、增加人工复核入口、完善权限体系和审计日志导出功能。对于本科毕业设计而言，当前系统已经形成可展示原型，后续重点是稳定性和规范性完善。",
    ],
}


def add_expansion(doc: Document, key: str) -> None:
    for text in EXPANSION_TEXT[key]:
        add_body(doc, text)


def add_chapter_one(doc: Document) -> None:
    add_chapter(doc, "第一章 绪论")
    add_body(doc, "本章首先说明课题的实现背景和意义，随后概述国内外相关技术实现现状，最后说明本文围绕基于 LangGraph 的多 Agent 协作框架所完成的主要工程实现工作。")
    add_section(doc, "1.1 实现背景和意义")
    for text in [
        "随着医院信息化、个人健康档案和体检服务的普及，普通用户能够获得的医疗资料越来越多。这些资料既包括体检总检报告、检验单、影像检查结论，也包括病历摘要、随访记录和健康管理建议。资料数量的增加并没有自动带来更清晰的健康信息组织，相反，不同机构的报告模板、指标命名、单位表达和异常标记存在差异，用户在长期管理资料时往往只能依赖人工翻阅，难以快速判断同一指标在不同时间点的变化，也难以知道某个提示结论来自哪一份原始报告。",
        "生成式大语言模型在医疗问答、文本总结和临床知识组织方面表现出较强的语言理解能力，国内中文文献已经讨论了其在医学场景中的机遇、限制和安全边界[9-13]。但医疗场景具有高敏感性和高责任属性，仅依赖一次性生成式回答无法满足审计要求。系统必须清楚记录数据来源、处理步骤、证据绑定、异常依据和人工复核边界，使每一条提示都能回到具体文档和具体字段，而不是停留在无法解释的自然语言结论。",
        "国家医院信息化建设和电子病历应用评价均强调数据标准化、质量控制、过程留痕和信息安全[1-2]，个人信息安全规范也要求在个人敏感信息处理过程中落实最小必要、访问控制和可追溯措施[3]。因此，本课题并不把系统实现目标设定为替代医生诊断，而是定位为医疗资料整理和医疗审计辅助：通过工程化方式把原始资料转化为结构化数据，再使用状态机式多 Agent 协作框架完成质量检查、风险提示、证据补全和报告输出。",
        "从工程实现角度看，传统线性工作流可以完成固定步骤的数据处理，但难以表达“发现证据不足后回到证据节点”“安全审查不通过后回到报告生成节点”“引用检查不通过后回到审计节点”等循环过程。LangGraph 提供了状态图、条件边和持久化状态的实现基础，适合把医疗审计过程组织为可回环、可观测、可恢复的状态机。本文围绕这一特点完成系统设计，使课题题目中的“基于 LangGraph 的多 Agent 协作框架”在系统核心流程中得到明确体现。",
    ]:
        add_body(doc, text)
    add_body(doc, "图1-1给出了本文系统面向的医疗审计场景与系统边界。图中将输入资料、处理流程、审计输出和非诊断性边界分开表示，强调系统输出是健康资料审计和风险提示，而不是临床诊断。")
    add_figure(doc, "图1-1_医疗审计场景与系统边界.png", "图1-1 医疗审计场景与系统边界", width_cm=14.6)

    add_section(doc, "1.2 国内外实现现状")
    for text in [
        "在医疗人工智能方向，国内研究主要围绕医疗大语言模型的机会、典型应用、知识增强方法、落地风险和创新应用展开，相关文献普遍认为医疗大模型不能脱离真实业务流程、数据治理和责任边界单独使用[9-13]。同时，电子病历文本挖掘、中文电子病历命名实体识别和真实世界病历数据治理等研究，为本文将原始医疗资料转化为结构化字段提供了直接参考[4-8]。",
        "在工程架构方向，检索增强生成、推理与工具调用结合、多 Agent 协作成为提升生成式系统可控性的常用方式。RAG 通过外部知识或业务数据补充模型上下文，降低脱离事实回答的概率[19]；ReAct 将推理轨迹和行动调用交替组织，使系统能够根据环境反馈调整任务计划[20]；AutoGen 等多 Agent 框架则强调用不同角色承担拆解、执行、评审和修正职责[21]。这些工作为本文拆分审计节点、保存中间状态和设计回环复核提供了参考。",
        "在软件工程实现方向，Web 系统需要稳定的数据模型、清晰的接口边界和可测试的服务层。软件工程和数据库教材为系统分层、接口边界、数据建模和事务处理提供了基础依据[14-15]；自然语言处理、机器学习和深度学习教材为 OCR 文本处理、指标抽取和语言模型能力边界提供了基本方法参考[16-18]。本文系统在后端使用这些工程化原则构建资料处理和审计接口，在智能能力部分通过 Provider 抽象将 OCR、标准化和 LLM 调用从业务流程中隔离出来，从而降低替换外部服务时对业务代码的影响。",
        "与单纯聊天机器人或固定工作流不同，本文实现重点在于把医疗资料处理、结构化入库和多 Agent 审计串成闭环。系统不仅要回答“生成了什么报告”，还要说明“报告从哪些文档来、经过哪些节点、每个节点输出什么、为什么可以结束或为什么需要回退”。这也是开题报告中“多 Agent 协作、任务审计、来源可追溯和前端展示”的延续。后续章节将围绕数据模型、接口设计和 LangGraph 状态机实现展开。",
    ]:
        add_body(doc, text)

    add_section(doc, "1.3 本文主要工作")
    for text in [
        "本文完成的主要工作包括四个方面。第一，设计并实现医疗资料数据底座，覆盖用户、健康记录、原始文件、OCR 结果、结构化文档、文档版本、指标、后台任务、任务事件、Provider 调用事件和审计报告运行记录等对象。所有关键对象均有明确字段和关联关系，能够支持后续论文中的 ER 图、字段表和接口说明。",
        "第二，设计并实现资料处理流程。系统支持文件上传、文本型 OCR 抽取、LLM/规则混合标准化、指标入库、版本快照和指标查询。标准化结果不直接覆盖原文，而是同时保留 OCR 原文、标准化 payload、版本 hash 和结构化指标，从数据结构上支持回溯和复核。",
        "第三，设计并实现基于 LangGraph 的综合审计报告生成流程。该流程不是一条直线，而是包含 audit_router 和 final_router 两个路由节点，并允许 Agent 节点完成后回到 audit_router；当引用检查或安全审查未通过时，final_router 能够回到审计或报告生成节点。该设计体现了状态机、条件边、回环复核和过程持久化。",
        "第四，完成前端模块和测试数据准备。前端包括文档接入、文档库、综合审计报告、智能洞察、RAG 知识库、指标探索和任务状态监控；Mock 数据包括多份体检、化验、病历、影像和报告型文档，用于支撑端到端测试、真实演示截图与论文图表展示。本文使用真实源码、数据库模型、运行页面和测试脚本生成图表，避免图表与项目实现脱节。",
    ]:
        add_body(doc, text)
    add_expansion(doc, "chapter_one")


def add_chapter_two(doc: Document, snapshot: dict) -> None:
    add_chapter(doc, "第二章 相关技术与系统基础")
    add_body(doc, "本章不把参考文献简单堆放在章末，而是围绕医疗资料标准化、医疗大模型与知识增强、Agent 工作流和 Provider 工程抽象四条线索，逐篇梳理已有研究的做法，并说明这些研究如何影响本文系统设计。表2-1列出本章核心引用文献及其对本文实现的直接启示。")
    literature_rows = [
        ["吴宗友等[4]", "综述电子病历文本挖掘，将任务归纳为清洗、集成、命名实体识别、关系抽取、检索和问答等环节。", "本文采用 OCR、标准化、指标入库、审计报告的分层链路。"],
        ["韩普等[6]", "提出结合外部语义特征、深度学习与 CRF/Voting 的中文电子病历实体识别方法，实验 F 值达到 94.06%。", "本文承认专门模型对标准语料有效，但本系统先做工程化标准化和证据留痕。"],
        ["崔少国等[7]", "融合汉字图像、五笔、医学词典和 Lattice 结构，提升中文电子病历命名实体识别效果。", "本文保留指标别名、单位和原文证据，避免只依赖通用模型。"],
        ["盖彦蓉等[8]", "真实世界电子病历抽取实验显示，模型在真实数据上的表现明显低于公开数据集，问题来自表述不规范、数据稀疏和治理不足。", "本文把版本、原文、任务事件和审计事件都保存下来，强调数据治理。"],
        ["Lewis 等[19]", "提出 RAG，将参数化模型和非参数化检索记忆结合，用外部文档提高知识密集任务的事实性和可解释性。", "本文把 RAG 作为审计状态机中的知识检索节点，而不是独立聊天功能。"],
        ["Yao 等[20]", "提出 ReAct，将推理轨迹和外部行动交替组织，使模型能够通过工具或知识库获得外部信息。", "本文把审计拆成节点和工具调用，避免一次性生成报告。"],
        ["Wu 等[21]", "AutoGen 将多 Agent 对话用于复杂任务，支持 LLM、工具和人类输入组合。", "本文借鉴职责拆分思想，但采用状态机而非纯聊天式 Agent。"],
        ["LangChain[22]", "LangGraph 强调长运行、有状态、可持久化的 Agent 工作流，并提供状态图、边和条件路由。", "本文用 LangGraph 表达审计回环、引用检查和安全审查。"],
    ]
    add_table(doc, "表2-1 相关研究引用依据与本文启示", ["文献", "已有工作做法", "本文设计启示"], literature_rows)
    add_section(doc, "2.1 医疗资料标准化与电子病历文本挖掘")
    for text in [
        "医疗资料标准化首先面对的是非结构化和半结构化文本。国家卫生健康委员会发布的医院信息化建设标准强调医疗数据、系统功能和质量管理的规范化要求[1]，电子病历应用水平分级评价也把数据质量、业务闭环和过程留痕作为医院信息化建设的重要内容[2]。个人信息安全规范进一步要求敏感个人信息处理具备明确目的、访问控制和可追溯能力[3]。因此，面向医疗审计的系统不能只保存一段模型摘要，而要把原始文件、识别文本、结构化字段、版本快照和最终证据之间的关系保存下来。",
        "吴宗友等在《电子病历文本挖掘研究综述》中将电子病历文本挖掘整理为数据清洗、数据集成、命名实体识别、关系抽取、信息检索和临床问答等任务，并指出电子病历文本包含大量非结构化信息，需要通过文本挖掘转化为可利用知识[4]。这一综述给本文的直接启示是：医疗资料审计不应从“报告生成”开始，而应先完成资料接入、OCR、标准化和指标索引，再进入审计报告状态机。",
        "杜晋华等围绕中文电子病历命名实体识别进行综述，关注疾病、症状、检查、药物、身体部位等实体类型，并指出中文病历文本存在术语密集、表达不规范和标注语料不足等问题[5]。韩普等提出融合外部语义特征、深度学习模型和 CRF 输出的中文电子病历实体识别方法，实验 F 值达到 94.06%，说明专门模型在标注数据和任务边界清晰时能够取得较好效果[6]。崔少国等进一步融合汉字图像、五笔编码、医学词典和 Lattice 结构，用语义与边界信息提升中文电子病历 NER 效果[7]。",
        "上述研究说明，中文医疗文本处理不是简单提示词可以完全解决的问题。专门 NER 模型强调实体边界、领域词典和训练语料，但本文面对的是个人体检报告、化验单、影像结论和病历摘要混杂的演示数据，难点不只在实体识别精度，还包括报告模板差异、指标单位差异、文档版本留痕和最终审计证据绑定。因此，本文没有把医疗资料处理简化为单个 NER 模型，而是采用“原始文件、OCR 结果、当前文档投影、文档版本、结构化指标”五层数据结构。",
        "盖彦蓉等针对真实世界中文电子病历开展知识抽取和数据治理分析，发现即使采用 BERT+Bi-LSTM+CRF 等主流模型，真实世界数据上的实体和关系识别效果仍明显低于公开数据集，原因包括表述不规范、数据稀疏、科室术语差异、隐私保护与数据利用平衡不足、全流程治理缺失等[8]。这与本文的工程取舍一致：系统需要保存 OCR 原文、标准化 payload、版本 hash、结构化指标、任务事件和审计事件，让后续报告能够回到具体数据来源，而不是只追求一次抽取结果。",
    ]:
        add_body(doc, text)

    add_section(doc, "2.2 医疗大模型、知识增强与 RAG 方法")
    for text in [
        "肖仰华和徐一丹从知识容器、能力引擎和自治智能体三个角度讨论生成式语言模型在医疗领域的机会，同时也指出医疗应用存在准确性、解释性、风险控制等局限[9]。康砚澜等进一步提出知识增强医学语言模型，认为知识图谱和多模态数据融合可以提升模型的专业性、准确性和可信性[10]。颜见智等系统梳理了生成式大语言模型在医疗服务、医学研究和教育中的典型应用，也强调幻觉、隐私保护、伦理、结果可控性和可解释性等挑战[11]。",
        "陈晓红等在医疗大模型技术及应用研究中提出，医疗大模型需要覆盖基础层、模型层、应用层和公共模块，并建立评价指标、数据集、模型对齐和评测平台等体系；该研究也把数据安全、技术风险、落地挑战和伦理道德列为核心问题[12]。何剑虎等采用系统性文献回顾方法梳理医疗大语言模型最新进展，指出未来研究应同时关注技术创新和伦理规范[13]。这些文献共同说明，医疗大模型不能脱离数据治理、评测体系和责任边界单独落地。",
        "国际研究也给出类似提醒。Thirunavukarasu 等在 Nature Medicine 发表的综述指出，LLM 在医学问答、临床文档和医学教育等场景中具有潜力，但临床使用仍需重视偏差、可解释性、安全性和监管问题[23]。Wornow 等在 npj Digital Medicine 中讨论面向电子健康记录的 LLM 和基础模型，强调 EHR 数据具有异质性、缺失性和编码差异，许多模型评估任务并不能充分说明其对真实医疗系统的有效性[24]。因此，本文系统只把大模型用于标准化、智能洞察和报告生成增强，不把生成结果直接等同于诊断结论。",
        "知识增强是降低生成式系统脱离事实的重要方向。Lewis 等提出 Retrieval-Augmented Generation，将参数化 seq2seq 模型与非参数化检索记忆结合，生成前先从外部文档索引中检索相关内容，并指出这种方式有助于提供来源、更新知识和缓解纯参数模型的事实性问题[19]。对于医疗资料审计来说，RAG 的价值不只是让回答更丰富，而是让报告能够附带知识来源、命中片段和审计依据，从而降低“凭空解释”的风险。",
        "在当前实现中，RAG 并不是独立的聊天功能，而是服务于综合审计报告状态机。系统内置医学审计知识块和默沙东医学手册摘要知识块，通过 knowledge_retrieval_agent 根据风险点和异常指标构造查询，使用确定性的词法排序召回相关知识，再将 knowledge_context 写入 AuditGraphState。最终报告中的 rag_summary、knowledge_sources 和 knowledge_chunk 类型证据，用于展示知识命中情况和来源。该设计对应 Lewis 等提出的“检索外部知识再生成”的思想[19]，也回应了医疗大模型研究中对事实性、可追溯性和责任边界的要求[9-13,23-24]。",
    ]:
        add_body(doc, text)

    add_section(doc, "2.3 Agent 工作流与状态机编排")
    for text in [
        "Yao 等提出的 ReAct 方法把语言模型的推理轨迹和任务行动交替组织：推理轨迹用于诱导、跟踪和更新行动计划，行动则用于与知识库或环境交互并获得额外信息[20]。该研究在问答、事实验证和交互式决策任务中验证了推理与行动结合的价值，也指出纯链式思考容易出现事实幻觉和错误传播。本文综合审计报告模块借鉴这一思想：审计流程不能只让模型一次性生成报告，而应允许节点根据状态补充证据、查询知识和修正报告。",
        "Wu 等提出的 AutoGen 将 LLM 应用组织为多个可对话、可配置的 Agent，并允许 Agent 结合 LLM、工具和人工输入完成复杂任务[21]。AutoGen 的重点是多 Agent 对话和灵活会话模式，这对本文有启发，但医疗审计场景并不适合完全采用开放式对话。原因在于审计流程需要确定的节点职责、状态字段、事件记录和失败回退路径，否则前端无法解释“流程为什么走到这里”。",
        "LangGraph 提供了更贴近本文需求的状态图式 Agent 工作流。其官方文档强调，LangGraph 面向长运行、有状态的工作流或 Agent，支持持久化执行、人类参与、记忆、调试可观测和状态迁移[22]。这与本文综合审计报告的需求一致：审计节点可能因为证据不足回到 evidence_agent，可能因为引用检查失败回到 audit_router，也可能因为安全审查失败回到 report_composer。",
        "本文采用 LangGraph 的原因正来自上述需求。系统中的 document_quality_agent、timeline_builder、measurement_consistency_agent、risk_agent、knowledge_retrieval_agent、evidence_agent、conflict_agent、compliance_agent、report_composer、citation_checker 和 safety_reviewer 都围绕同一个 AuditGraphState 工作。每个节点只负责一个有限任务，路由节点根据 completed_agents、citation_issues、safety_issues、needs_report_revision 等字段决定下一步。相比 AutoGen 式的开放会话，本文更强调状态机可控性和审计可追溯性。",
        "图2-1说明了本文使用 LangGraph 的基本机制。START 进入状态加载节点后，audit_router 按状态分派多个审计 Agent；普通 Agent 完成后回到 audit_router；报告链路完成后进入 final_router；final_router 根据引用检查和安全审查结果决定回到 audit_router、回到 report_composer 或进入 persist_report。该机制使系统具备状态机和回环能力，也使前端可以根据事件流展示节点高亮和边流转。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图2-2_LangGraph状态机机制.png", "图2-1 LangGraph 状态机机制", width_cm=14.8)

    add_section(doc, "2.4 Provider 抽象与外部能力接入")
    providers = snapshot["providers"]
    ocr_config = providers.get("ocr", "auto")
    rows = [
        ["OCRProvider", ocr_config, "负责从文件字节中抽取文本；当前配置为 auto，文本样本走 plaintext，图片样本在密钥配置后走 baidu_ocr accurate。"],
        ["NormalizationProvider", providers.get("normalization", "llm_direct"), "负责把 OCR 原文转换为文档类型、报告日期、指标数组和叙事事实；当前为 llm_direct，并带规则兜底。"],
        ["LLMProvider", providers.get("llm", "openai_compatible"), "负责对话、洞察、报告生成增强和部分标准化能力，当前通过 openai_compatible 接入兼容模型服务。"],
        ["StorageProvider", providers.get("storage", "database_inline"), "负责文件存储，当前使用 database_inline，便于本地演示和测试复现。"],
    ]
    add_table(doc, "表2-2 Provider 抽象及当前配置", ["抽象接口", "当前配置", "功能说明"], rows)
    for text in [
        "Provider 抽象的核心作用是隔离业务流程与外部能力。OCR、LLM 和存储服务具有明显的不稳定性和环境差异：本地演示时可能只处理纯文本，生产部署时可能需要百度 OCR、视觉模型或对象存储；标准化能力也可能在规则解析、LLM 解析和混合解析之间切换。如果业务服务直接依赖某个具体 SDK，后续替换成本会很高，也不利于记录调用耗时、错误类别和重试结果。该设计符合软件工程中通过接口边界降低模块耦合、提高可维护性的基本思想[14]。",
        "图2-2展示了本文系统的 Provider 接入方式。业务服务只依赖统一接口，ProviderGateway 负责记录 provider_events、映射异常、统计耗时并隐藏外部服务差异。论文中的测试环境读取 .env 后得到当前配置：ocr_provider 为 auto，normalization_provider 为 llm_direct，llm_provider 为 openai_compatible，storage_provider 为 database_inline。该图由项目配置和 Provider 代码生成，反映当前真实工程状态。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图2-1_Provider抽象与外部能力接入.png", "图2-2 Provider 抽象与外部能力接入", width_cm=14.8)

    add_section(doc, "2.5 本项目已有实现基础")
    for text in [
        "本文实现不是从空白项目开始，而是在开题报告确定的工程路线基础上完成系统化整理和强化。已有基础包括 FastAPI 后端、SQLAlchemy 模型、MySQL 数据库连接、用户认证、文件上传、OCR 结果保存、标准化服务、指标查询、智能洞察对话、RAG 知识库、综合审计报告接口、React 前端页面和 Mock 体检数据脚本。中期阶段暴露出的主要问题是系统展示没有充分体现 LangGraph 架构，因此本文后续实现重点转向综合审计报告模块和状态机可视化。",
        "在当前版本中，后端路由统一挂载到 /api，已包含 auth、files、ocr、ingestion、documents、document-versions、records、measurements、query、tasks、knowledge、insight、chat 和 audit-reports 等分组。数据库模型覆盖核心业务对象，符合关系数据库通过主键、外键和关系模式组织业务对象的建模思路[15]；测试脚本能够创建 admin@qq.com 测试账户并写入 25 份 Mock 文档，其中包含 20 份体检类单据、5 份报告型单据和 85 条结构化指标。综合审计报告模块已经能够持久化 audit_report_runs、audit_report_events 和 audit_report_node_states。",
        "因此，本文后续章节的图表和表格均以当前源码为依据生成。需要特别说明的是，当前实现中的审计节点以规则审计、知识检索和证据绑定为主，LLM 主要用于标准化、智能洞察和报告生成增强；本文不会虚假描述为每个节点都由大模型自主生成，而是准确表述为“LangGraph 编排多个审计节点，结合规则审计、RAG 检索、证据绑定、条件路由和 LLM 增强能力”。这种表述更符合项目真实状态，也更符合工程类毕业设计的要求。",
    ]:
        add_body(doc, text)
    add_expansion(doc, "chapter_two")


def add_chapter_three(doc: Document, snapshot: dict) -> None:
    add_chapter(doc, "第三章 需求分析与总体设计")
    add_body(doc, "本章从功能需求、总体架构、业务流程、后端接口、数据库结构和安全追溯等方面进行总体设计，说明系统如何将文档处理、结构化入库和综合审计报告生成组织为完整工程闭环。")
    add_section(doc, "3.1 需求分析")
    rows = [
        ["资料接入", "用户能够上传体检报告、检验单、病历摘要和影像文本结论，系统保存原始文件和文件元数据。", "record_files"],
        ["OCR 与标准化", "系统能够抽取文本并生成结构化文档、版本快照和指标记录。", "ocr_results、extracted_documents、document_versions、measurements"],
        ["指标查询", "用户能够按名称、时间和文档查询结构化指标，用于后续审计和趋势展示。", "measurements API"],
        ["综合审计", "系统能够选择多个文档版本，通过 LangGraph 多节点流程生成综合审计报告。", "audit_report_runs、events、node_states"],
        ["过程可视化", "前端能够展示审计节点、边流转、最终报告和历史记录。", "audit-reports API"],
        ["安全边界", "系统输出为非诊断性风险提示，所有结论必须保留来源或提示人工复核。", "evidence_items、safety_reviewer"],
    ]
    add_table(doc, "表3-1 系统功能需求与数据支撑", ["需求", "说明", "支撑对象"], rows)
    for text in [
        "系统的用户角色为个人健康资料管理者。用户使用系统的基本过程为：登录账号，上传资料，触发 OCR 或文本抽取，执行标准化入库，在文档库中查看文档版本和指标，在综合审计报告模块中选择若干报告生成审计报告。由于本课题面向医疗审计辅助而非诊断，系统必须在输出层明确边界，避免出现“确诊”“治愈”“必须立即用药”等不适合审计场景的表述。",
        "非功能需求主要包括可追溯性、可扩展性、稳定性和可测试性。可追溯性要求关键数据从原始文件到最终报告都有路径；可扩展性要求 OCR、标准化和 LLM 服务可替换；稳定性要求外部服务失败时有规则兜底或明确错误状态；可测试性要求接口、数据库和 LangGraph 审计流程能在 Mock 数据下复现。本文设计的数据库、Provider 抽象和 audit_report_events 正是围绕这些非功能需求展开。",
    ]:
        add_body(doc, text)

    add_section(doc, "3.2 系统总体架构")
    for text in [
        "系统采用前后端分离和分层后端架构。前端负责用户交互、文件选择、模块切换、数据展示和审计图流转；后端提供统一 REST 接口，按照 API 层、服务层、Provider 层和数据访问层组织；数据库保存业务数据、任务事件和审计状态；外部能力通过 Provider 层接入。综合审计报告模块位于服务层核心位置，负责把文档版本和指标数据组织为 LangGraph 初始状态，再将每个节点输出写入审计事件表。",
        "图3-1展示系统总体架构。与开题报告中的技术路线一致，该架构从资料接入开始，经过 OCR、结构化抽取、版本化存储和指标查询，最终进入多 Agent 审计与前端展示。图中把数据库和 Provider 能力放在后端服务的支撑层，说明系统不是单一前端页面，也不是单一大模型调用，而是一个包含数据闭环、服务编排和审计留痕的工程系统。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图3-1_系统总体架构.png", "图3-1 系统总体架构", width_cm=14.8)

    add_section(doc, "3.3 业务流程设计")
    for text in [
        "业务流程围绕“上传资料后形成可审计数据资产”展开。用户上传文件后，系统创建 records 和 record_files，随后通过 OCR 接口生成 ocr_results。标准化接口读取 OCR 原文，将文档类型、报告日期、结构化指标和叙事事实写入 extracted_documents，同时创建 document_versions 和 measurements。综合审计报告模块再读取选中的文档版本和指标，启动 LangGraph 状态机，持续产生事件和节点状态，最后持久化 final_report。",
        "图3-2给出了医疗资料处理与审计业务流程。流程图强调两个闭环：第一个闭环是从原始文件到结构化指标的入库闭环，第二个闭环是从文档版本到综合审计报告的状态机闭环。前者保证系统有真实数据可审计，后者保证审计过程不是黑盒生成。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图3-2_医疗资料处理与审计业务流程.png", "图3-2 医疗资料处理与审计业务流程", width_cm=14.8)

    add_section(doc, "3.4 后端分层与接口设计")
    route_rows = []
    for name in [
        "auth",
        "files",
        "ocr",
        "ingestion",
        "documents",
        "document-versions",
        "measurements",
        "tasks",
        "knowledge",
        "audit-reports",
        "insight",
        "chat",
    ]:
        routes = snapshot["routes"].get(name, [])
        route_rows.append([name, str(len(routes)), "；".join(routes[:3]) + ("；..." if len(routes) > 3 else "")])
    add_table(doc, "表3-2 主要 API 分组", ["接口分组", "数量", "典型接口"], route_rows)
    for text in [
        "后端分层结构如图3-3所示。API 层负责认证、参数解析和响应模型；服务层负责业务编排，例如文件上传、OCR 任务、标准化入库和综合审计报告；Provider 层封装外部能力；模型层由 SQLAlchemy ORM 定义数据库结构。这样的分层避免了接口函数直接操作复杂业务流程，也使论文中的每个功能点都能定位到具体代码层。",
        "API 设计遵循资源化思路。files 分组负责文件上传和文件关联数据查询；ocr 分组负责 OCR 修订记录；ingestion 分组负责将 OCR 结果标准化；documents 和 document-versions 分组负责文档与版本查询；measurements 分组负责指标搜索和时序；tasks 分组负责后台任务状态、事件和 provider-events/summary 监控；knowledge 分组负责知识块、来源和检索；audit-reports 分组负责综合审计报告创建、执行、事件轮询和节点状态查询。表3-2中的接口数量来自当前 FastAPI APIRouter 快照。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图3-3_后端分层结构.png", "图3-3 后端分层结构", width_cm=14.8)
    add_figure(doc, "图3-5_API接口分组图.png", "图3-5 API 接口分组图", width_cm=14.8)

    add_section(doc, "3.5 数据库设计")
    core_rows = [
        ["users", "用户账号", ", ".join(snapshot["tables"]["users"])],
        ["record_files", "原始文件与文件字节", ", ".join(snapshot["tables"]["record_files"])],
        ["ocr_results", "OCR 识别结果与修订", ", ".join(snapshot["tables"]["ocr_results"])],
        ["extracted_documents", "当前标准化文档投影", ", ".join(snapshot["tables"]["extracted_documents"])],
        ["document_versions", "标准化文档版本快照", ", ".join(snapshot["tables"]["document_versions"])],
        ["measurements", "结构化指标索引", ", ".join(snapshot["tables"]["measurements"])],
        ["audit_report_runs", "综合审计报告运行", ", ".join(snapshot["tables"]["audit_report_runs"][:8]) + "..."],
    ]
    add_table(doc, "表3-3 核心数据表职责", ["数据表", "职责", "主要字段"], core_rows)
    for text in [
        "数据库设计是系统可追溯能力的基础。本文将资料处理主链路、异步任务链路和审计报告状态机链路分开建模。资料处理主链路从 users 到 records、record_files、ocr_results、extracted_documents、document_versions 和 measurements；任务链路由 tasks、task_events 和 provider_events 记录；审计链路由 audit_report_runs、audit_report_events 和 audit_report_node_states 记录。三条链路通过 user_id、record_id、document_version_id、task_id 和 run_id 建立关联。",
        "图3-4为核心数据库 ER 图。图中字段来自 SQLAlchemy 模型快照，保留主键、外键和关键业务字段。由于 tasks 和 audit_report_runs 字段较多，图中展示论文分析所需字段，完整字段清单放在附录 A。该设计能够支持从最终报告向前追溯到审计事件、节点输出、文档版本、OCR 原文和原始文件。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图3-4_核心数据库ER图.png", "图3-4 核心数据库 ER 图", width_cm=15.0)

    add_section(doc, "3.6 安全与可追溯设计")
    for text in [
        "医疗资料通常包含敏感个人信息，系统设计必须避免无边界扩散。本文在数据层通过 user_id 进行用户隔离，业务查询均围绕当前用户过滤；在文件和 OCR 层保留原文与 provider 信息；在任务层记录 task_events 和 provider_events；在审计层记录 audit_report_events 和 audit_report_node_states；在输出层通过 safety_reviewer 检查不适合审计场景的医疗措辞。",
        "图3-6展示安全与可追溯设计。其重点不是做复杂权限体系，而是在毕业设计原型系统范围内保证每次资料处理、Provider 调用和审计报告生成都有可查记录。对于医疗审计辅助系统而言，可追溯性本身也是安全边界的一部分，因为无法追溯来源的提示不应进入最终报告。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图3-6_安全与可追溯设计.png", "图3-6 安全与可追溯设计", width_cm=14.8)
    add_expansion(doc, "chapter_three")


def add_chapter_four(doc: Document) -> None:
    add_chapter(doc, "第四章 系统详细实现")
    add_body(doc, "本章说明文件上传与 OCR、标准化入库、LangGraph 审计状态机、审计事件持久化和前端模块的具体实现，重点体现系统中真实数据如何在后端状态图和前端流程图之间保持一致。")
    add_section(doc, "4.1 文件上传与 OCR 处理实现")
    for text in [
        "文件上传流程由前端提交文件，后端 files API 校验并写入 records 和 record_files。record_files 不仅保存 original_filename 和 display_name，也保存 content_type、size_bytes、storage_provider、storage_key 和 content_bytes。当前演示环境采用 database_inline 存储，便于本地复现和测试；若后续接入对象存储，只需要替换 StorageProvider。",
        "OCR 处理通过 ocr API 触发。系统创建或更新 tasks，并将具体抽取动作交给 OCRProvider。当前 Provider 配置为 auto：文本型 Mock 报告路由到 plaintext，图片型资料在百度 OCR 密钥配置后路由到 baidu_ocr accurate，后续仍保留 openai_compatible_vision 扩展点。处理完成后，系统写入 ocr_results，字段包括 record_file_id、revision_number、supersedes_ocr_result_id、is_current、provider_name、status、raw_text、raw_payload 和 created_at。",
        "图4-1展示文件上传与 OCR 处理时序。该图从前端、files API、FileUploadService、数据库、ocr API、TaskProcessor 和 OCRProvider 之间的调用关系出发，说明系统把长耗时任务和外部能力调用从同步页面操作中拆开。即使本地演示仍以文本样本为主，任务和事件结构也已经为真实 OCR 服务、失败重试和状态监控保留了工程扩展空间。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图4-1_文件上传与OCR处理时序.png", "图4-1 文件上传与 OCR 处理时序", width_cm=14.8)

    add_section(doc, "4.2 标准化与版本化入库实现")
    for text in [
        "标准化服务读取 OCRResult.raw_text 后调用 NormalizationProvider。当前实现为 llm_direct，并带规则兜底。Provider 输出包含 document_type、document_category、report_date、measurements 和 prose_facts 等字段。系统不会只保存模型输出文本，而是将结构化结果写入 ExtractedDocument、DocumentVersion 和 Measurement。这样既能支持前端文档展示，也能支持指标搜索和 LangGraph 审计。",
        "标准化稳定性的关键在于 schema 约束、规则兜底和版本机制。LLM 输出容易受到提示词、输入格式和模型状态影响，因此系统需要将结果约束为可解析 payload；当 LLM 调用失败或字段缺失时，规则解析可以至少抽取日期、常见指标和原文事实；当同一文档重新标准化时，DocumentVersion.snapshot_hash 用于避免重复版本或定位差异。该设计降低了“标准化效果不稳定”对后续审计流程的影响。",
        "图4-2展示标准化与版本化入库流程。OCR 原文进入 NormalizationProvider，输出 NormalizationResult 后形成 ExtractedDocument 当前投影、DocumentVersion 版本快照和 Measurement 指标索引。图中还标注了稳定性策略：LLM 失败时规则兜底，snapshot_hash 避免重复版本，narrative_context 不强行生成数值指标。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图4-2_标准化与版本化入库流程.png", "图4-2 标准化与版本化入库流程", width_cm=14.8)

    add_section(doc, "4.3 综合审计报告 LangGraph 实现")
    node_rows = [
        ["load_graph_state", "初始化状态", "读取选中文档版本、指标、运行配置", "next_action、route_history"],
        ["audit_router", "主路由", "根据 completed_agents 和证据状态选择下一个审计节点", "next_action"],
        ["document_quality_agent", "文档质量审计", "检查是否缺少文档、OCR 原文、报告日期", "quality_findings"],
        ["timeline_builder", "时间线构建", "按报告日期和指标时间生成时间线", "timeline"],
        ["measurement_consistency_agent", "指标一致性审计", "按规则检查血糖、ALT、CRP、白细胞等异常", "consistency_findings"],
        ["risk_agent", "风险提示", "将异常指标转化为非诊断性关注项", "risk_findings"],
        ["knowledge_retrieval_agent", "知识检索", "根据异常指标、风险点和医学审计知识块执行 RAG 召回", "knowledge_queries、knowledge_context"],
        ["evidence_agent", "证据绑定", "为风险、冲突和指标结论绑定原文或指标证据", "evidence_items"],
        ["conflict_agent", "冲突复核", "检查叙事事实与结构化指标是否存在复核点", "conflict_findings"],
        ["compliance_agent", "合规检查", "检查关键结论是否仍缺少证据", "compliance_findings"],
        ["quality_gate", "质量门禁", "判断是否具备进入报告生成的最低条件", "quality_gate"],
        ["report_composer", "报告生成", "组织最终报告结构、结论和证据项", "report_draft"],
        ["citation_checker", "引用检查", "检查非 info 结论是否绑定证据", "citation_issues"],
        ["safety_reviewer", "安全审查", "拦截不适合审计场景的医疗表述", "safety_issues"],
        ["final_router", "终态路由", "根据引用和安全检查决定回退或持久化", "next_action、stop_reason"],
        ["persist_report", "持久化", "写入最终报告并结束图执行", "final_report"],
    ]
    add_table(doc, "表4-1 LangGraph 节点职责与输入输出", ["节点", "职责", "输入依据", "主要输出"], node_rows)
    for text in [
        "综合审计报告模块是本文体现 LangGraph 架构的核心。服务层创建 AuditReportRun 后，把用户选择的 document_version_ids 转换为 AuditGraphState。状态中包含 documents、measurements、knowledge_chunks、knowledge_context、completed_agents、route_history、iteration_count、max_iterations、quality_findings、risk_findings、evidence_items、citation_issues、safety_issues 和 final_report 等字段。每个节点只返回局部更新，Engine 将更新合并回状态。",
        "状态机的第一阶段由 audit_router 控制。audit_router 首先检查文档质量，再构建时间线，随后进行指标一致性审计、风险提示、知识检索和证据绑定。当风险、冲突或知识来源缺少证据时，_findings_need_evidence 会使路由回到 evidence_agent。普通审计节点完成后均回到 audit_router，这种回环不是前端动画，而是 LangGraph 中真实存在的条件边和状态迁移。",
        "状态机的第二阶段由 final_router 控制。report_composer 生成报告草稿后，citation_checker 检查关键结论是否绑定证据，safety_reviewer 检查报告中是否包含不适合审计场景的医疗承诺。如果 citation_issues 存在，final_router 会回到 audit_router 补充审计；如果 safety_issues 存在，则回到 report_composer 重写报告；只有检查通过或达到最大迭代保护条件时，流程才进入 persist_report。",
        "图4-3是综合审计报告 LangGraph 状态流转图，边关系来自 AUDIT_GRAPH_EDGES。图中蓝线表示 audit_router 分派审计节点，灰线表示节点完成后回到 audit_router，橙线表示报告生成与审查链路，红线表示 final_router 回环补充审计，绿线表示通过后持久化并结束。该图直接对应前端综合报告模块中需要展示的真实流转结构。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图4-3_综合审计报告LangGraph状态流转图.png", "图4-3 综合审计报告 LangGraph 状态流转图", width_cm=15.2)

    add_section(doc, "4.4 审计事件与节点状态持久化")
    for text in [
        "为了让前端看到“数据正在节点上流动”的真实效果，后端不能只在执行结束后返回报告，而需要在执行过程中记录事件。AuditGraphEngine.stream 每产生一个节点输出，服务层就保存一条 AuditReportEvent，并更新 AuditReportNodeState。事件包含 sequence、event_type、node_name、edge_source、edge_target、status、message 和 payload；节点状态包含 node_name、status、visit_count、last_event_id 和 output。",
        "这种持久化方式使前端可以通过 GET /audit-reports/{run_id}/events 和 GET /audit-reports/{run_id}/nodes 轮询。前端不需要猜测流程走到哪里，而是根据真实事件高亮节点和边。若某个节点被回环再次访问，visit_count 会增加，route_history 也会体现重复流转。对于论文而言，这一点能够证明系统不是静态流程图，而是有后端状态支撑的 LangGraph 运行展示。",
        "图4-4展示审计事件与节点状态持久化关系。AuditReportRun 保存整体运行状态和最终报告，AuditReportEvent 保存边和节点执行事件，AuditReportNodeState 保存每个节点的最新状态和输出，前端通过轮询接口展示节点高亮、边流转和完整报告按钮。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图4-4_审计事件与节点状态持久化.png", "图4-4 审计事件与节点状态持久化", width_cm=14.8)

    add_section(doc, "4.5 前端模块实现")
    for text in [
        "前端以成熟产品工作台形式组织模块，主要包括文档接入、文档库、综合审计报告、智能洞察、RAG 知识库和指标探索。文档接入负责上传和处理资料；文档库展示标准化文档、文档类型和创建时间；指标探索支持结构化指标查询；智能洞察保持聊天机器人形态并提供历史会话列表；RAG 知识库展示知识来源、检索命中和 BM25 词法匹配结果；综合审计报告模块独立承担文档选择、LangGraph 流程图流转和最终报告展示。",
        "前端设计的重点是把智能洞察、RAG 知识库和综合审计报告分离。智能洞察是对话式能力，保留历史消息和会话列表；综合审计报告不是聊天框，而是一个以流程图为核心的工作台。用户选择若干文档后启动报告生成，流程图按照后端事件流实时高亮节点和边，最终通过按钮打开完整报告。右上角任务状态按钮用于查看 OCR 与标准化任务状态，避免长耗时任务在页面上变成不可解释的等待。",
        "图4-5展示前端模块结构。该图来自 frontend/src/App.jsx 的模块配置和 API 封装设计，说明各模块共享 Token 和 API 服务，但在交互目标上保持分工。综合审计报告模块是论文中最能体现课题题目的前端入口。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图4-5_前端模块结构.png", "图4-5 前端模块结构", width_cm=14.8)
    add_expansion(doc, "chapter_four")


def add_chapter_five(doc: Document, snapshot: dict) -> None:
    add_chapter(doc, "第五章 测试与运行效果分析")
    add_body(doc, "本章基于 Mock 体检数据、核心接口测试、端到端运行链路和真实演示截图说明系统运行效果，并对当前版本的完成度、可展示内容和后续完善方向进行分析。")
    add_section(doc, "5.1 Mock 体检数据构造")
    rows = [
        ["演示账号", "admin@qq.com / 123123123", "来自 scripts/seed_admin_mock_data.py"],
        ["文档总数", "25", "20 份体检类单据，5 份报告型单据"],
        ["结构化指标", "85", "写入 measurements，用于指标查询和审计"],
        ["文档类型", "体检、检验、影像、病历摘要、报告型文档", "覆盖文档库、RAG 和综合审计演示"],
        ["知识块", "17", "9 条默沙东医学手册摘要知识块，来自 scripts/seed_msd_manual_rag.py"],
    ]
    add_table(doc, "表5-1 Mock 数据构成", ["项目", "数量/分布", "来源说明"], rows)
    for text in [
        "为了支撑完整流程测试，项目提供 seed_admin_mock_data.py 脚本创建 admin@qq.com 测试账户，并写入 25 份模拟医疗资料。该数据不是只为页面凑数，而是覆盖体检报告、化验报告、病历摘要、影像结论和综合报告型文档，使系统能够同时验证结构化指标、叙事事实、文档版本、RAG 知识检索和综合审计报告。",
        "Mock 数据的价值在于让系统能够稳定复现端到端流程。文本样本可以稳定进入 OCR plaintext 分支，图片样本可在百度 OCR 密钥配置后验证真实识别链路；标准化结果写入 ExtractedDocument、DocumentVersion 和 Measurement；综合审计报告可以选择多份文档运行 LangGraph，并在报告中展示文档证据和知识来源。图5-1展示 Mock 体检数据分布。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图5-1_Mock体检数据分布.png", "图5-1 Mock 体检数据分布", width_cm=14.8)

    add_section(doc, "5.2 接口与端到端测试")
    test_rows = [
        ["T01", "用户登录", "POST /auth/login", "返回 access_token", "通过"],
        ["T02", "文件上传", "POST /files/upload", "创建 records 和 record_files", "通过"],
        ["T03", "OCR 抽取", "POST /ocr/files/{id}/extract", "创建 ocr_results 和任务事件", "通过"],
        ["T04", "标准化入库", "POST /ingestion/ocr-results/{id}/normalize", "创建文档、版本和指标", "通过"],
        ["T05", "文档查询", "GET /documents", "返回用户文档列表", "通过"],
        ["T06", "指标搜索", "GET /measurements/search", "返回匹配指标", "通过"],
        ["T07", "任务状态监控", "GET /tasks/provider-events/summary", "返回 OCR 与标准化 Provider 调用概览", "通过"],
        ["T08", "知识库检索", "GET /knowledge/search", "返回知识块命中和来源", "通过"],
        ["T09", "综合审计创建", "POST /audit-reports", "创建 audit_report_runs", "通过"],
        ["T10", "综合审计执行", "POST /audit-reports/{run_id}/execute", "产生事件、节点状态和最终报告", "通过"],
        ["T11", "事件轮询", "GET /audit-reports/{run_id}/events", "返回边和节点执行事件", "通过"],
        ["T12", "节点轮询", "GET /audit-reports/{run_id}/nodes", "返回节点状态和 visit_count", "通过"],
    ]
    add_table(doc, "表5-2 端到端测试用例", ["编号", "功能", "接口/对象", "预期结果", "结果"], test_rows)
    for text in [
        "测试围绕从登录到报告生成的主链路设计。首先验证用户认证是否能够返回 Token，然后验证文件上传、OCR、标准化、文档查询和指标查询，再验证任务状态监控、知识库检索、综合审计报告创建、执行、事件轮询和节点状态轮询。该测试路径覆盖了系统的核心业务数据流，也覆盖了论文中各图表对应的主要模块。",
        "项目当前后端健康检查接口返回 200，前端本地服务能够登录 admin@qq.com 并读取 Mock 数据。需要说明的是，毕业设计阶段的测试以功能闭环验证为主，后续上线前还应补充更多异常场景，例如空文档、缺失日期、标准化失败、LLM Provider 超时、审计报告安全审查不通过和最大迭代次数达到上限等。",
        "图5-2展示端到端测试链路。该链路从登录和上传开始，经过 OCR、标准化、查询、审计运行、节点轮询和最终报告，能够对应用户演示时的完整操作路径。与纯后端单元测试相比，该链路更适合作为毕业答辩演示主线。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图5-2_端到端测试链路.png", "图5-2 端到端测试链路", width_cm=14.8)

    add_section(doc, "5.3 真实演示运行截图")
    for text in [
        "为了避免运行效果只停留在流程图和表格层面，本文在本地启动后端与前端服务，使用 admin@qq.com 演示账号登录系统，并截取真实页面作为运行证据。截图对应当前代码和 Mock 数据状态，能够展示工作台、任务队列、文档库、RAG 知识库、综合审计流程和完整报告预览。",
    ]:
        add_body(doc, text)
    add_screenshot(doc, "01_workspace_home.png", "图5-3 系统工作台与功能模块", width_cm=15.0)
    add_screenshot(doc, "02_task_queue_monitor.png", "图5-4 任务队列与 OCR/标准化状态监控", width_cm=15.0)
    add_screenshot(doc, "03_document_vault.png", "图5-5 文档库中的标准化文档", width_cm=15.0)
    add_screenshot(doc, "04_rag_knowledge_search.png", "图5-6 RAG 知识库检索与来源展示", width_cm=15.0)
    add_screenshot(doc, "06_audit_graph_completed.png", "图5-7 综合审计报告 LangGraph 运行结果", width_cm=15.0)
    add_screenshot(doc, "07_audit_report_modal.png", "图5-8 综合审计报告正文与 RAG 来源证据", width_cm=15.0)

    add_section(doc, "5.4 运行效果与不足分析")
    for text in [
        "从功能完成度看，当前项目已经具备后端演示和论文撰写的核心条件。资料上传、OCR 保存、标准化入库、文档库、指标查询、任务监控、RAG 知识库、LangGraph 综合审计报告、事件持久化和前端模块均已形成闭环。对于毕业设计而言，最重要的是能够展示“为什么必须使用 LangGraph”：系统中的审计流程包含多个审计节点、两个路由节点、状态字段、条件边和回环复核，而不是简单的一条线性流水线。",
        "从技术体现看，综合审计报告模块解决了中期阶段“没有体现 LangGraph 架构”的问题。前端流程图高亮来自后端 audit_report_events 和 audit_report_node_states，后端事件来自 AuditGraphEngine.stream，边关系来自 AUDIT_GRAPH_EDGES。用户在页面上看到的不是静态流程图，而是真实运行事件驱动的节点状态。RAG 命中知识块后，报告预览还能展示 knowledge_sources 和 knowledge_chunk 证据。",
        "从论文定稿角度看，测试部分还应强调测试数据与系统功能之间的对应关系。Mock 数据中的体检报告主要用于验证总检结论、时间线和多个指标并存的情况；化验报告主要用于验证数值指标、单位和异常标记；病历摘要主要用于验证叙事事实和用药、既往史等自然语言内容；报告型文档用于验证较长文本能否参与审计和摘要。",
        "对于接口测试结果，论文中不应只列出接口是否返回 200，还应说明接口通过后对数据库产生了什么影响。例如上传接口通过后应能在 record_files 中看到文件记录，OCR 接口通过后应能在 ocr_results 和 task_events 中看到处理结果，标准化接口通过后应能看到 document_versions 和 measurements，综合审计执行接口通过后应能看到 audit_report_events 中连续递增的 sequence。",
        "从不足看，当前系统仍有改进空间。第一，OCR Provider 虽已具备 auto 路由和百度 OCR 接入能力，但图片、扫描 PDF 和复杂表格的识别质量仍需要更多真实样本验证。第二，标准化效果虽然加入 llm_direct 和规则兜底，但仍需继续完善指标别名、单位换算和参考范围解析。第三，审计节点当前以规则审计、RAG 检索和证据绑定为主，后续可以在安全边界明确的前提下引入更多 LLM 解释型节点。第四，数据库迁移链和任务队列在上线前仍需加强，保证部署环境和本地环境一致。",
        "性能方面，当前系统更关注流程可解释和演示稳定，而不是高并发。由于毕业设计场景主要面向单用户或少量用户演示，系统把文件内容直接保存到数据库中可以降低部署复杂度。若后续进入真实生产环境，应将文件迁移到对象存储，并为 OCR、标准化和审计报告生成任务增加更完整的队列、重试、限流和失败恢复机制。",
        "综上，测试与运行效果章节不仅证明系统能跑通，还证明系统为什么这样设计。Mock 数据证明输入足够多样，接口测试证明后端链路完整，事件与节点状态证明 LangGraph 不是静态概念，RAG 检索证明知识来源可以进入报告，前端截图证明状态机可以被用户理解，最终报告证明审计结果可以落地阅读。",
    ]:
        add_body(doc, text)


def add_conclusion(doc: Document) -> None:
    add_chapter(doc, "结束语")
    for text in [
        "本文围绕《基于 LangGraph 的多 Agent 协作框架在医疗审计场景的设计与实现》完成了系统需求分析、总体架构设计、数据库设计、接口设计、核心流程实现和测试说明。系统以个人健康资料为入口，完成文件上传、OCR、标准化入库、版本化存储、指标查询、RAG 知识检索、任务状态监控、综合审计报告和前端展示，形成了从原始资料到审计报告的工程闭环。",
        "本文实现的关键特点在于使用 LangGraph 将综合审计报告生成过程组织为状态机。audit_router 负责分派文档质量、时间线、指标一致性、风险、知识检索、证据、冲突、合规和质量门禁等节点；final_router 负责根据引用检查和安全审查决定回退或持久化；事件和节点状态被保存到数据库并用于前端高亮展示。该设计使系统能够体现多 Agent 协作、条件路由、回环复核和可追溯执行过程。",
        "后续工作主要包括三点。第一，持续验证真实 OCR 服务和更多真实报告样本，提升图片、PDF 和复杂表格的处理能力。第二，完善标准化规则库、指标单位换算和知识库内容，增强不同报告模板之间的一致性。第三，在保持非诊断性边界的前提下扩展 LLM 审计节点，使风险解释、证据摘要和用户可读报告更加自然。总体来看，当前系统已经能够支撑毕业设计论文初稿和后续答辩演示，后续重点是补充测试、优化细节和完善论文表达。",
    ]:
        add_body(doc, text)


def add_references(doc: Document) -> None:
    add_chapter(doc, "参考文献")
    refs = [
        "国家卫生健康委员会规划与信息司. 全国医院信息化建设标准与规范（试行）[S/OL]. 2018[2026-05-07]. https://www.nhc.gov.cn/ewebeditor/uploadfile/2018/04/20180413162542120.pdf.",
        "国家卫生健康委办公厅. 关于印发电子病历系统应用水平分级评价管理办法（试行）及评价标准（试行）的通知: 国卫办医函〔2018〕1079号[Z/OL]. 2018[2026-05-07]. https://www.nhc.gov.cn/yzygj/c100068/201812/b01f63185ef74a41afa30adeb5c58ccf.shtml.",
        "国家市场监督管理总局, 国家标准化管理委员会. 信息安全技术 个人信息安全规范: GB/T 35273-2020[S]. 北京: 中国标准出版社, 2020.",
        "吴宗友, 白昆龙, 杨林蕊, 等. 电子病历文本挖掘研究综述[J]. 计算机研究与发展, 2021,58(3):513-527.",
        "杜晋华, 尹浩, 冯嵩. 中文电子病历命名实体识别的研究与进展[J]. 电子学报, 2022,50(12):3030-3053.",
        "韩普, 刘亦卓, 李晓艳. 基于深度学习和多特征融合的中文电子病历实体识别研究[J]. 南京大学学报(自然科学), 2019,55(6):942-951.",
        "崔少国, 陈俊桦, 李晓虹. 融合语义及边界信息的中文电子病历命名实体识别[J]. 电子科技大学学报, 2022,51(4):565-571.",
        "盖彦蓉, 张云秋, 张慧, 等. 面向知识抽取的真实世界中文电子病历数据质量分析与治理对策研究[J]. 医学信息学杂志, 2025,46(12):47-53.",
        "肖仰华, 徐一丹. 大规模生成式语言模型在医疗领域的应用: 机遇与挑战[J]. 医学信息学杂志, 2023,44(9):1-11.",
        "康砚澜, 郭倩宇, 张文强, 等. 基于知识增强的医学语言模型: 现状、技术与应用[J]. 医学信息学杂志, 2023,44(9):12-22.",
        "颜见智, 何雨鑫, 骆子烨, 等. 生成式大语言模型在医疗领域的潜在典型应用与面临的挑战[J]. 医学信息学杂志, 2023,44(9):23-31.",
        "陈晓红, 刘浏, 袁依格, 等. 医疗大模型技术及应用发展研究[J]. 中国工程科学, 2024,26(6):77-88.",
        "何剑虎, 王德健, 赵志锐, 等. 大语言模型在医疗领域的前沿研究与创新应用[J]. 医学信息学杂志, 2024,45(9):10-18.",
        "张海藩, 牟永敏. 软件工程导论[M]. 6版. 北京: 清华大学出版社, 2013.",
        "王珊, 萨师煊. 数据库系统概论[M]. 5版. 北京: 高等教育出版社, 2014.",
        "宗成庆. 统计自然语言处理[M]. 2版. 北京: 清华大学出版社, 2013.",
        "周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "邱锡鹏. 神经网络与深度学习[M]. 北京: 机械工业出版社, 2020.",
        "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems. 2020,33:9459-9474.",
        "Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]//International Conference on Learning Representations. 2023.",
        "Wu Q, Bansal G, Zhang J, et al. AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation[EB/OL]. arXiv:2308.08155, 2023[2026-05-07]. https://arxiv.org/abs/2308.08155.",
        "LangChain. LangGraph overview[EB/OL]. [2026-05-07]. https://docs.langchain.com/oss/python/langgraph/overview.",
        "Thirunavukarasu A J, Ting D S J, Elangovan K, et al. Large language models in medicine[J]. Nature Medicine, 2023,29:1930-1940. DOI:10.1038/s41591-023-02448-8.",
        "Wornow M, Xu Y, Thapa R, et al. The shaky foundations of large language models and foundation models for electronic health records[J]. npj Digital Medicine, 2023,6:135. DOI:10.1038/s41746-023-00879-8.",
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False)
        run = p.add_run("[")
        set_run_font(run, 10.5, name="Times New Roman")
        add_bookmarked_reference_number(p, idx, size=10.5)
        run = p.add_run(f"] {ref}")
        set_run_font(run, 10.5)


def add_acknowledgement(doc: Document) -> None:
    add_chapter(doc, "致谢")
    for text in [
        "本课题从开题到系统实现过程中，指导教师在题目定位、工程路线、文档规范和阶段检查方面给予了持续指导。特别是在中期检查后，课题重新聚焦到 LangGraph 多 Agent 协作框架本身，使系统从一般健康资料管理工具转向更切题的医疗审计状态机实现。",
        "感谢学院提供毕业设计规范、论文模板、参考文献著录说明和进度节点安排，使论文撰写能够按照统一格式推进。感谢开源社区提供 FastAPI、SQLAlchemy、Pydantic、LangGraph、React 和相关工具链，为系统实现提供了基础能力。",
    ]:
        add_body(doc, text)


def add_appendices(doc: Document, snapshot: dict) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_chapter(doc, "附录 A 核心数据表字段清单")
    for table_name, fields in snapshot["tables"].items():
        add_subsection(doc, f"A.{list(snapshot['tables']).index(table_name) + 1} {table_name}")
        add_body(doc, "字段名：" + "，".join(fields), first_line=False)

    add_chapter(doc, "附录 B 论文插图来源对照表")
    rows = [
        ["图1-1", "scripts/generate_thesis_assets.py", "项目业务边界与开题报告技术路线"],
        ["图2-1", "app/core/config.py、.env、Provider 代码", "当前 Provider 配置和外部能力抽象"],
        ["图2-2", "app/services/audit_graph/engine.py", "LangGraph 状态机机制"],
        ["图3-1", "app/main.py、app/api/router.py、frontend 模块", "系统总体架构"],
        ["图3-2", "文件、OCR、标准化和审计服务源码", "医疗资料处理与审计流程"],
        ["图3-3", "app/api、app/services、app/providers、app/models", "后端分层结构"],
        ["图3-4", "SQLAlchemy Base.metadata", "核心数据库 ER 图"],
        ["图3-5", "FastAPI APIRouter 快照", "API 接口分组"],
        ["图3-6", "用户隔离、任务事件、审计事件模型", "安全与可追溯设计"],
        ["图4-1", "files API、ocr API、TaskProcessor、OCRProvider", "文件上传与 OCR 时序"],
        ["图4-2", "NormalizationProvider、DocumentVersion、Measurement", "标准化与版本化入库流程"],
        ["图4-3", "AUDIT_GRAPH_EDGES", "综合审计报告 LangGraph 状态流转图"],
        ["图4-4", "audit_report_runs/events/node_states", "审计事件与节点状态持久化"],
        ["图4-5", "frontend/src/App.jsx 模块配置", "前端模块结构"],
        ["图5-1", "scripts/seed_mock_exam_data.py", "Mock 体检数据分布"],
        ["图5-2", "测试链路与 pytest/health 检查结果", "端到端测试链路"],
    ]
    add_table(doc, "表B-1 论文插图来源对照表", ["图号", "来源", "说明"], rows)


def count_text_chars(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    return len(re.sub(r"\s+", "", text))


def main() -> None:
    snapshot = load_snapshot()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(sec, header_text="摘要", page_start=1, page_fmt="upperRoman")
    add_chinese_abstract(doc)
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(sec, header_text="Abstract", page_start=2, page_fmt="upperRoman")
    add_english_abstract(doc)
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(sec, header_text="目录", page_start=3, page_fmt="upperRoman")
    add_toc(doc)
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(sec, header_text=TITLE, page_start=1, page_fmt="decimal")
    add_chapter_one(doc)
    add_chapter_two(doc, snapshot)
    add_chapter_three(doc, snapshot)
    add_chapter_four(doc)
    add_chapter_five(doc, snapshot)
    add_conclusion(doc)
    add_references(doc)
    add_acknowledgement(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    output = OUTPUT
    try:
        doc.save(str(output))
    except PermissionError:
        output = FALLBACK_OUTPUT
        doc.save(str(output))
    print(output.relative_to(ROOT))
    print(f"text_chars={count_text_chars(output)}")


if __name__ == "__main__":
    main()
